#!/usr/bin/env python3
"""Generate CFO Briefing Word document for Roxster, Ltd. - March 2026"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.space_before = Pt(0)

# ─── Helper Functions ───

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, font_size=9, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color

def add_section_banner(doc, text, font_size=11):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(4)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1F3864')
    pPr.append(shd)
    run = para.add_run(f'  {text}')
    run.bold = True
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return para

def add_blue_subheader(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return para

def add_body_text(doc, text, bold=False, font_size=9.5):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Calibri'
    run.bold = bold
    return para

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'B0B0B0')
        borders.append(el)
    tblPr.append(borders)

def fmt_currency(val, decimals=0):
    if decimals == 0:
        return f"${val:,.0f}"
    return f"${val:,.{decimals}f}"

def fmt_pct(val):
    return f"{val:.1f}%"

# ─── 1. CONFIDENTIAL BANNER ───
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
para.paragraph_format.space_after = Pt(2)
run = para.add_run('CONFIDENTIAL — FOR INTERNAL USE ONLY')
run.bold = True
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ─── 2. TITLE BANNER ───
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
para.paragraph_format.space_before = Pt(2)
para.paragraph_format.space_after = Pt(4)
pPr = para._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), '1F3864')
pPr.append(shd)
run = para.add_run('  F & W Properties Inc.  |  Roxster, Ltd.  |  March 2026 Monthly Review')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0xFF, 0xF2, 0xCC)

# ─── 3. MEMO HEADER TABLE ───
memo_table = doc.add_table(rows=4, cols=2)
memo_table.alignment = WD_TABLE_ALIGNMENT.LEFT
memo_data = [
    ('TO:', 'Mr. Soul, Chief Financial Officer'),
    ('FROM:', 'Property Finance Team'),
    ('DATE:', 'March 17, 2026'),
    ('RE:', 'Monthly P&L, Valuation & Action Items — Roxbury Arms (1571-1619 Roxbury Road, Grandview Heights, OH 43212)')
]
for i, (label, value) in enumerate(memo_data):
    set_cell_text(memo_table.cell(i, 0), label, bold=True, font_size=9)
    set_cell_text(memo_table.cell(i, 1), value, font_size=9)
    memo_table.cell(i, 0).width = Inches(0.8)
    memo_table.cell(i, 1).width = Inches(5.7)

doc.add_paragraph()  # spacer

# ─── 4. EXECUTIVE SUMMARY ───
add_section_banner(doc, 'EXECUTIVE SUMMARY')

add_blue_subheader(doc, 'Key Performance Indicators — March 2026 Budget')

# KPI boxes as a table (1 row, 5 cols)
kpi_table = doc.add_table(rows=2, cols=5)
kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(kpi_table)

kpi_labels = ['Budget NOI\n(Monthly)', 'Total Operating\nIncome (Monthly)', 'Budgeted Cash Flow\n(Monthly)', 'Est. Property\nValue', 'Est. Equity\nResidual']
kpi_values = ['$53,485', '$110,238', '$33,263', '$8.38M', '$2.93M']

for j in range(5):
    cell = kpi_table.cell(0, j)
    set_cell_shading(cell, 'D9E2F3')
    set_cell_text(cell, kpi_labels[j], bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell = kpi_table.cell(1, j)
    set_cell_shading(cell, 'D9E2F3')
    set_cell_text(cell, kpi_values[j], bold=True, font_size=11, color=RGBColor(0x1F, 0x38, 0x64), align=WD_ALIGN_PARAGRAPH.CENTER)

add_body_text(doc, (
    'Roxster, Ltd. enters March 2026 with a budgeted monthly NOI of $53,485 (Jan-Jun), '
    'reflecting total operating income of $110,238 against operating expenses of $56,754. '
    'Year-over-year, budgeted operating income of $1,335,287 represents a 2.2% increase over '
    '2025 actual income of $1,307,035. After debt service and a $341,000 capital expenditure program, '
    'the property is projected to generate $44,987 in annual cash flow after CapEx. '
    'The income-approach valuation at a 7.5% cap rate yields an estimated property value of $8.38M '
    'and equity residual of $2.93M after closing costs and debt payoff.'
))

# ─── 5. FINANCIAL PERFORMANCE — MARCH 2026 ───
add_section_banner(doc, 'FINANCIAL PERFORMANCE — MARCH 2026')

add_blue_subheader(doc, 'Income Analysis')
add_body_text(doc, (
    'Budgeted total operating income for March 2026 is $110,238 (Jan-Jun rate), composed of '
    'rents ($105,153), fee income ($863), utility reimbursements ($3,450), and other income ($772). '
    'This compares to March 2025 actual income of $102,835, a 7.2% YoY increase, and March 2024 '
    'actual income of $104,144, a 5.9% increase over two years.'
))

add_blue_subheader(doc, 'Expense Analysis')
add_body_text(doc, (
    'Budgeted total operating expenses for March 2026 are $56,754 (Jan-Jun rate). Key line items include '
    'repairs & maintenance ($14,043/mo), payroll ($23,012/mo), utilities ($6,174/mo), property tax ($3,454/mo), '
    'property insurance ($3,658/mo), and professional services ($6,295/mo). Note: March 2025 expenses were '
    'anomalously low at $47,470 due to payroll of only $6,091, well below the 2026 budget of $23,012. '
    'March 2024 expenses were $61,977, driven by a payroll spike of $28,958.'
))

add_blue_subheader(doc, 'P&L Comparison — March (Monthly)')

# P&L comparison table
pl_table = doc.add_table(rows=12, cols=4)
pl_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(pl_table)

headers = ['Line Item', 'Mar 2024 Actual', 'Mar 2025 Actual', 'Mar 2026 Budget']
for j, h in enumerate(headers):
    cell = pl_table.cell(0, j)
    set_cell_shading(cell, '1F3864')
    set_cell_text(cell, h, bold=True, font_size=8.5, color=RGBColor(0xFF, 0xFF, 0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)

pl_data = [
    ('Total Operating Income', '$104,144', '$102,835', '$110,238'),
    ('  Rents', '—', '—', '$105,153'),
    ('  Fee Income', '—', '—', '$863'),
    ('  Utility Reimbursement', '—', '—', '$3,450'),
    ('  Other Income', '—', '—', '$772'),
    ('Total Operating Expense', '$61,977', '$47,470', '$56,754'),
    ('  Payroll', '$28,958', '$6,091', '$23,012'),
    ('  Repairs & Maintenance', '$12,037', '$22,182', '$14,043'),
    ('  Taxes & Insurance', '$4,936', '$4,793', '$7,112'),
    ('  Utilities', '$10,495', '$7,112', '$6,174'),
    ('Net Operating Income', '$42,167', '$55,365', '$53,485'),
]

for i, row_data in enumerate(pl_data):
    for j, val in enumerate(row_data):
        r = i + 1
        cell = pl_table.cell(r, j)
        is_total = row_data[0].startswith('Total') or row_data[0].startswith('Net')
        bold = is_total
        align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_text(cell, val, bold=bold, font_size=8.5, align=align)
        if is_total:
            set_cell_shading(cell, 'F2F2F2')

# ─── 6. PROPERTY VALUATION — INCOME APPROACH ───
add_section_banner(doc, 'PROPERTY VALUATION — INCOME APPROACH')

add_blue_subheader(doc, 'Valuation Key Metrics')

val_table = doc.add_table(rows=2, cols=4)
val_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(val_table)

val_labels = ['Adjusted NOI', 'Cap Rate', 'Est. Property Value', 'Est. Debt Outstanding']
val_values = ['$628,654', '7.5%', '$8,382,055', '$5,200,000']
for j in range(4):
    cell = val_table.cell(0, j)
    set_cell_shading(cell, 'D9E2F3')
    set_cell_text(cell, val_labels[j], bold=True, font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell = val_table.cell(1, j)
    set_cell_shading(cell, 'D9E2F3')
    set_cell_text(cell, val_values[j], bold=True, font_size=10, color=RGBColor(0x1F, 0x38, 0x64), align=WD_ALIGN_PARAGRAPH.CENTER)

add_body_text(doc, (
    f'Closing costs at 3% = $251,462. Estimated equity residual after debt payoff = $2,930,593.'
))

add_blue_subheader(doc, 'Sensitivity Analysis — Equity Residual by NOI Scenario & Cap Rate')

# Sensitivity matrix
noi_scenarios = [
    ('Stress -10%', 565789),
    ('Base -5%', 597221),
    ('Budget', 628654),
    ('Upside +5%', 660087),
    ('Upside +10%', 691519),
]
cap_rates = [0.060, 0.065, 0.070, 0.075, 0.080, 0.085]

sens_table = doc.add_table(rows=len(cap_rates) + 1, cols=len(noi_scenarios) + 1)
sens_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(sens_table)

# Header row
cell = sens_table.cell(0, 0)
set_cell_shading(cell, '1F3864')
set_cell_text(cell, 'Cap Rate \\ NOI', bold=True, font_size=8, color=RGBColor(0xFF, 0xFF, 0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
for j, (label, noi) in enumerate(noi_scenarios):
    cell = sens_table.cell(0, j + 1)
    set_cell_shading(cell, '1F3864')
    set_cell_text(cell, f'{label}\n${noi:,.0f}', bold=True, font_size=7.5, color=RGBColor(0xFF, 0xFF, 0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)

# Data rows
for i, cr in enumerate(cap_rates):
    cell = sens_table.cell(i + 1, 0)
    set_cell_shading(cell, 'D9E2F3')
    set_cell_text(cell, f'{cr:.1%}', bold=True, font_size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for j, (label, noi) in enumerate(noi_scenarios):
        equity = (noi / cr) * (1 - 0.03) - 5200000
        cell = sens_table.cell(i + 1, j + 1)
        # Color coding
        if equity < 0:
            bg = 'FFC7CE'
        elif label == 'Budget' and abs(cr - 0.075) < 0.001:
            bg = 'FFEB9C'
        elif equity > 3500000:
            bg = 'C6EFCE'
        elif equity > 2500000:
            bg = 'FFEB9C'
        else:
            bg = 'FFFFFF'
        # Budget/base cap rate highlight
        if label == 'Budget' and abs(cr - 0.075) < 0.001:
            bg = 'FFEB9C'
        set_cell_shading(cell, bg)
        set_cell_text(cell, fmt_currency(equity), font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT,
                      bold=(label == 'Budget' and abs(cr - 0.075) < 0.001))

add_body_text(doc, (
    'The sensitivity matrix illustrates equity residual outcomes across NOI stress/upside scenarios '
    'and cap rate assumptions. The base case (Budget NOI at 7.5% cap) yields $2.93M equity residual. '
    'At a compressed 6.0% cap rate with upside NOI, residual could exceed $5.9M. Conversely, '
    'stress NOI at an 8.5% cap rate compresses residual to approximately $1.3M.'
), font_size=9)

# ─── 7. 2026 ANNUAL CASH FLOW OUTLOOK ───
add_section_banner(doc, '2026 ANNUAL CASH FLOW OUTLOOK')

add_blue_subheader(doc, 'Monthly Cash Flow Projection')

cf_table = doc.add_table(rows=9, cols=3)
cf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(cf_table)

cf_headers = ['Line Item', 'Jan-Jun (Monthly)', 'Jul-Dec (Monthly)']
for j, h in enumerate(cf_headers):
    cell = cf_table.cell(0, j)
    set_cell_shading(cell, '1F3864')
    set_cell_text(cell, h, bold=True, font_size=8.5, color=RGBColor(0xFF, 0xFF, 0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)

cf_data = [
    ('Total Operating Income', '$110,238', '$112,309'),
    ('Total Operating Expense', '($56,754)', '($56,857)'),
    ('Net Operating Income', '$53,485', '$55,452'),
    ('Mortgage Interest', '($18,316)', '($18,316)'),
    ('Principal Reduction', '($1,907)', '($1,907)'),
    ('Cash Flow Before CapEx', '$33,263', '$35,230'),
    ('Capital Expenditures (Annual)', '', '$341,000'),
    ('Projected Cash Flow After CapEx (Annual)', '', '$44,987'),
]

for i, row_data in enumerate(cf_data):
    for j, val in enumerate(row_data):
        r = i + 1
        cell = cf_table.cell(r, j)
        is_bold_row = 'Net Operating' in row_data[0] or 'Cash Flow' in row_data[0]
        align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_text(cell, val, bold=is_bold_row, font_size=8.5, align=align)
        if 'After CapEx' in row_data[0]:
            set_cell_shading(cell, 'D9E2F3')
        elif is_bold_row:
            set_cell_shading(cell, 'F2F2F2')

add_blue_subheader(doc, 'Annual Summary')

ann_table = doc.add_table(rows=2, cols=5)
ann_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(ann_table)

ann_labels = ['Annual Operating\nIncome', 'Annual Operating\nExpense (Adj.)', 'Annual NOI\n(Adjusted)', 'Annual Debt\nService', 'Annual Cash Flow\nAfter CapEx']
ann_values = ['$1,335,287', '$706,633', '$628,654', '$242,667', '$44,987']
for j in range(5):
    cell = ann_table.cell(0, j)
    set_cell_shading(cell, '1F3864')
    set_cell_text(cell, ann_labels[j], bold=True, font_size=7.5, color=RGBColor(0xFF, 0xFF, 0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
    cell = ann_table.cell(1, j)
    set_cell_shading(cell, 'D9E2F3')
    set_cell_text(cell, ann_values[j], bold=True, font_size=9.5, color=RGBColor(0x1F, 0x38, 0x64), align=WD_ALIGN_PARAGRAPH.CENTER)

add_body_text(doc, (
    'The 2026 budget anticipates $341,000 in capital expenditures concentrated in H1 (Jan-Jun: $305K), '
    'with the largest outlays in January ($75K), February ($75K), and June ($95K). CapEx tapers to zero '
    'by Q4, improving monthly cash flow in the second half of the year. After all debt service and CapEx, '
    'projected annual cash flow is $44,987.'
))

# ─── 8. MD&A ───
add_section_banner(doc, "MANAGEMENT DISCUSSION & ANALYSIS (MD&A)")

add_blue_subheader(doc, 'Revenue Analysis')
add_body_text(doc, (
    'Budgeted 2026 operating income of $1,335,287 reflects a 2.2% increase over 2025 actuals ($1,307,035). '
    'The primary growth driver is a scheduled rent increase effective July 2026, lifting monthly rents from '
    '$105,153 to $107,224 (+$2,071/mo, +2.0%). Fee income and utility reimbursements remain stable. '
    'The property continues to benefit from strong occupancy in the Grandview Heights submarket.'
))

add_blue_subheader(doc, 'Expense Analysis')
add_body_text(doc, (
    'Adjusted operating expenses of $706,633 represent a 3.2% increase over 2025 actual expenses ($684,638). '
    'The adjusted figure includes a $24,968 net adjustment ($36,057 admin increase offset by -$11,089 payroll '
    'reduction). Repairs & maintenance at $14,043/mo is budgeted between the 2025 anomalous high of $22,182 '
    'and the 2024 level of $12,037, reflecting a return to normalized maintenance spend. '
    'Payroll normalizes at $23,012/mo after the 2025 anomaly ($6,091 in March 2025) and 2024 spike ($28,958).'
))

add_blue_subheader(doc, 'Operational Highlights')
bullets = [
    'Rent increase scheduled July 2026: +$2,071/mo (+2.0%) across portfolio.',
    'Capital program front-loaded: 89% of CapEx ($305K of $341K) deployed in H1 2026.',
    'Payroll budget normalized at $23,012/mo; management should monitor for variance given 2024-2025 volatility.',
    'Property tax and insurance budgeted at $7,112/mo combined, consistent with recent actuals.',
]
for b in bullets:
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_after = Pt(1)
    for run in para.runs:
        run.font.size = Pt(9)
    # Clear and reset
    para.clear()
    run = para.add_run(b)
    run.font.size = Pt(9)
    run.font.name = 'Calibri'

add_blue_subheader(doc, 'Risks & Opportunities')

risk_table = doc.add_table(rows=5, cols=2)
risk_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(risk_table)

for j, h in enumerate(['Risks', 'Opportunities']):
    cell = risk_table.cell(0, j)
    set_cell_shading(cell, '1F3864')
    set_cell_text(cell, h, bold=True, font_size=8.5, color=RGBColor(0xFF, 0xFF, 0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)

risks = [
    'Payroll volatility: 2024-2025 swings suggest staffing/cost control risk.',
    'R&M overrun: 2025 March actual ($22,182) significantly exceeded budget ($14,043).',
    'CapEx front-loading creates H1 cash flow pressure.',
    'Rising cap rates could compress property value and equity residual.',
]
opps = [
    'Rent growth: July increase provides immediate NOI uplift in H2.',
    'CapEx completion: zero CapEx in Q4 creates strong free cash flow.',
    'Cap rate compression: at 6.5%, equity residual rises to ~$4.2M.',
    'Utility reimbursement optimization could offset rising utility costs.',
]
for i in range(4):
    set_cell_text(risk_table.cell(i + 1, 0), risks[i], font_size=8.5)
    set_cell_text(risk_table.cell(i + 1, 1), opps[i], font_size=8.5)

add_blue_subheader(doc, 'Outlook')
add_body_text(doc, (
    'The 2026 budget positions Roxbury Arms for stable performance with modest income growth and controlled '
    'expenses. The front-loaded capital program addresses deferred maintenance while preserving H2 cash flow. '
    'Management should closely monitor payroll variance given historical volatility and track R&M actuals '
    'against the $14,043/mo budget. The July rent increase provides a natural inflection point for NOI '
    'improvement. Overall, the property is projected to generate positive cash flow of $44,987 after all '
    'obligations, supporting continued equity growth.'
))

# ─── 9. ACTION ITEMS FOR CFO REVIEW ───
add_section_banner(doc, 'ACTION ITEMS FOR CFO REVIEW')

action_table = doc.add_table(rows=6, cols=4)
action_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(action_table)

action_headers = ['#', 'Action Item', 'Owner', 'Target Date']
for j, h in enumerate(action_headers):
    cell = action_table.cell(0, j)
    set_cell_shading(cell, '1F3864')
    set_cell_text(cell, h, bold=True, font_size=8.5, color=RGBColor(0xFF, 0xFF, 0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)

actions = [
    ('1', 'Review and approve 2026 operating budget and CapEx schedule', 'CFO', 'Mar 31, 2026'),
    ('2', 'Confirm July 2026 rent increase implementation plan', 'Property Mgmt', 'May 15, 2026'),
    ('3', 'Establish payroll variance monitoring protocol (monthly)', 'Controller', 'Apr 15, 2026'),
    ('4', 'Evaluate R&M spend vs. budget through Q1 actuals', 'Property Mgmt', 'Apr 30, 2026'),
    ('5', 'Assess refinancing options given current debt service load', 'CFO / Treasury', 'Jun 30, 2026'),
]
for i, (num, item, owner, date) in enumerate(actions):
    set_cell_text(action_table.cell(i + 1, 0), num, font_size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(action_table.cell(i + 1, 1), item, font_size=8.5)
    set_cell_text(action_table.cell(i + 1, 2), owner, font_size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(action_table.cell(i + 1, 3), date, font_size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)

# Set column widths for action table
action_table.cell(0, 0).width = Inches(0.4)
action_table.cell(0, 1).width = Inches(4.0)
action_table.cell(0, 2).width = Inches(1.2)
action_table.cell(0, 3).width = Inches(1.2)

doc.add_paragraph()  # spacer

add_blue_subheader(doc, 'CFO Sign-Off')

sign_table = doc.add_table(rows=3, cols=2)
sign_table.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_borders(sign_table)

sign_data = [
    ('Reviewed By:', ''),
    ('Signature:', ''),
    ('Date:', ''),
]
for i, (label, val) in enumerate(sign_data):
    set_cell_text(sign_table.cell(i, 0), label, bold=True, font_size=9)
    set_cell_text(sign_table.cell(i, 1), val, font_size=9)
    sign_table.cell(i, 0).width = Inches(1.2)
    sign_table.cell(i, 1).width = Inches(3.5)

# Footer
doc.add_paragraph()
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('Prepared by Property Finance Team  |  F & W Properties Inc.  |  March 17, 2026')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
run.italic = True

# Save
output_path = '/Users/natefisher/Projects/hatfield/cfo-agent/output/Roxster_CFO_Briefing_Mar2026.docx'
doc.save(output_path)
print(f'Document saved to {output_path}')
