#!/usr/bin/env python3
import os, json, re, subprocess, smtplib, glob
from datetime import datetime, timedelta
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from pathlib import Path
import anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIREFLIES_API_KEY = os.environ.get("FIREFLIES_API_KEY")
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY")
GMAIL_USER        = "mrsoulcfo@gmail.com"
GMAIL_APP_PW      = (os.environ.get("GMAIL_APP_PASSWORD") or "").replace("\xa0", "").replace(" ", "")
EMAIL_TO          = "nfisher@peak10group.com"
OUTPUT_DIR        = Path.home() / "meeting-sync" / "Meeting Summaries"
LOG_FILE          = Path.home() / "meeting-sync" / "sync.log"
HOURS_LOOKBACK    = 12

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

def log(msg):
    ts = datetime.now().strftime("%a %b %d %H:%M:%S %Y")
    line = f"[{ts}] {msg}"
    print(line)
    with open(LOG_FILE, "a") as f:
        f.write(line + "\n")

def fetch_recent_transcripts():
    cutoff_ms = int((datetime.now() - timedelta(hours=HOURS_LOOKBACK)).timestamp() * 1000)
    query = '{"query": "{ transcripts { id title date summary { overview } } }"}'
    result = subprocess.run([
        "curl", "-s", "-X", "POST", "https://api.fireflies.ai/graphql",
        "-H", "Content-Type: application/json",
        "-H", f"Authorization: Bearer {FIREFLIES_API_KEY}",
        "-d", query
    ], capture_output=True, text=True)
    data = json.loads(result.stdout)
    transcripts = data["data"]["transcripts"]
    recent = [t for t in transcripts if t["date"] >= cutoff_ms and t.get("summary")]
    log(f"Found {len(recent)} new transcript(s) in last {HOURS_LOOKBACK} hours")
    return recent

def generate_structured_summary(transcript):
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    title = transcript["title"]
    date  = datetime.fromtimestamp(transcript["date"]/1000).strftime("%B %d, %Y")
    raw   = transcript["summary"]["overview"] if transcript.get("summary") else ""

    prompt = f"""You are Hatfield, an executive assistant AI for Peak 10 Group.

Convert this meeting summary into structured JSON. Return ONLY valid JSON, no markdown:
{{
  "meeting": "{title}",
  "date": "{date}",
  "attendees": "Extract from summary or write Unknown",
  "prepared_by": "Hatfield (AI)",
  "sections": [
    {{
      "title": "SECTION TITLE",
      "bullets": ["bullet 1", "bullet 2"]
    }}
  ],
  "action_items": [
    {{
      "action": "Description",
      "owner": "Person",
      "due": "Date or status"
    }}
  ]
}}

Meeting title: {title}
Date: {date}
Summary:
{raw}

Create logical sections based on content. Keep bullets concise and factual."""

    message = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=2000,
        messages=[{"role": "user", "content": prompt}]
    )
    return json.loads(message.content[0].text.strip())

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def create_docx(data, output_path):
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"Meeting Notes — {data['date']}")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    doc.add_paragraph()

    htable = doc.add_table(rows=2, cols=4)
    htable.style = 'Table Grid'
    header_data = [
        [("Meeting", data["meeting"]), ("Date", data["date"])],
        [("Attendees", data["attendees"]), ("Prepared by", data["prepared_by"])],
    ]
    for r_idx, row_data in enumerate(header_data):
        row = htable.rows[r_idx]
        for c_idx, (label, value) in enumerate(row_data):
            label_cell = row.cells[c_idx * 2]
            value_cell = row.cells[c_idx * 2 + 1]
            set_cell_bg(label_cell, "1F497D")
            run = label_cell.paragraphs[0].add_run(label)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
            value_cell.paragraphs[0].add_run(value).font.size = Pt(10)

    doc.add_paragraph()

    for i, section in enumerate(data["sections"], 1):
        heading = doc.add_paragraph()
        run = heading.add_run(f"{i}. {section['title']}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        for bullet in section["bullets"]:
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(bullet).font.size = Pt(11)
        doc.add_paragraph()

    ai_heading = doc.add_paragraph()
    run = ai_heading.add_run("ACTION ITEMS SUMMARY")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    ai_table = doc.add_table(rows=1, cols=3)
    ai_table.style = 'Table Grid'
    for idx, (cell, hdr) in enumerate(zip(ai_table.rows[0].cells, ["Action Item", "Owner", "Due / Status"])):
        set_cell_bg(cell, "1F497D")
        run = cell.paragraphs[0].add_run(hdr)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    for item in data["action_items"]:
        row = ai_table.add_row()
        row.cells[0].paragraphs[0].add_run(item.get("action", "")).font.size = Pt(10)
        row.cells[1].paragraphs[0].add_run(item.get("owner",  "")).font.size = Pt(10)
        row.cells[2].paragraphs[0].add_run(item.get("due",    "")).font.size = Pt(10)

    for row in ai_table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = [Inches(4.0), Inches(1.75), Inches(1.75)][idx]

    doc.add_paragraph()
    footer = doc.add_paragraph()
    run = footer.add_run("Generated by Hatfield · Peak 10 Group AI")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(output_path)
    log(f"Saved: {output_path}")

def send_email(subject, body, attachment_path):
    msg = MIMEMultipart()
    msg["From"]    = GMAIL_USER
    msg["To"]      = EMAIL_TO
    msg["Subject"] = subject.encode("ascii", errors="replace").decode("ascii")
    msg.attach(MIMEText(body, "plain", "utf-8"))
    with open(attachment_path, "rb") as f:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(f.read())
    encoders.encode_base64(part)
    part.add_header("Content-Disposition", f'attachment; filename="{Path(attachment_path).name}"')
    msg.attach(part)
    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(GMAIL_USER, GMAIL_APP_PW)
        server.sendmail(GMAIL_USER, EMAIL_TO, msg.as_bytes())
    log(f"Email sent: {subject}")

def main():
    log("=== Meeting Summary Generator started ===")
    transcripts = fetch_recent_transcripts()
    if not transcripts:
        log("No new transcripts to process.")
        return
    for t in transcripts:
        try:
            title = t["title"].replace("\xa0", " ").encode("ascii", errors="replace").decode("ascii")
            date  = datetime.fromtimestamp(t["date"]/1000).strftime("%Y-%m-%d")
            safe  = re.sub(r'[^\w\s-]', '', title).strip().replace(" ", "_")
            fname = f"{date}_{safe}_Summary.docx"
            fpath = OUTPUT_DIR / fname
            log(f"Processing: {title}")
            data = generate_structured_summary(t)
            create_docx(data, fpath)
            subject = f"Meeting Summary: {title} ({date})"
            body    = f"Hi Nate,\n\nAttached is the meeting summary for:\n\n  {title}\n  {date}\n\nGenerated by Hatfield.\n\n— Mr. Soul CFO"
            send_email(subject, body, fpath)
        except Exception as e:
            log(f"ERROR processing '{t['title']}': {e}")
    log("=== Done ===")

if __name__ == "__main__":
    main()
