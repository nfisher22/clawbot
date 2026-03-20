#!/usr/bin/env python3
"""
Hatfield Fireflies Agent — pulls meeting transcripts and summaries via GraphQL
"""
import os
import requests
from datetime import datetime, timezone
from vault_secrets import get_secrets
get_secrets()

FIREFLIES_API_KEY = os.getenv("FIREFLIES_API_KEY")
FIREFLIES_URL = "https://api.fireflies.ai/graphql"
AUDIT_LOG = "/opt/clawbot/logs/audit.log"

def audit(level, script, message):
    ts = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    line = f"{ts} | {level} | {script} | {message}\n"
    try:
        with open(AUDIT_LOG, "a") as f:
            f.write(line)
    except Exception:
        pass

def graphql(query, variables=None):
    headers = {
        "Authorization": f"Bearer {FIREFLIES_API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {"query": query}
    if variables:
        payload["variables"] = variables
    response = requests.post(FIREFLIES_URL, headers=headers, json=payload, timeout=15)
    return response.json()

def get_recent_transcripts(limit=5):
    query = """
    query {
        transcripts(limit: %d) {
            id
            title
            date
            duration
            summary {
                overview
                action_items
            }
            organizer_email
        }
    }
    """ % limit
    try:
        data = graphql(query)
        transcripts = data.get("data", {}).get("transcripts", [])
        audit("INFO", "fireflies-agent", f"Retrieved {len(transcripts)} transcripts")
        return transcripts
    except Exception as e:
        audit("ERROR", "fireflies-agent", f"Exception: {e}")
        return []

def format_transcripts(transcripts):
    if not transcripts:
        return "No recent meetings found in Fireflies."
    lines = [f"🎙️ *Recent Meetings* ({len(transcripts)} found)\n"]
    for t in transcripts:
        title = t.get("title", "Untitled")
        raw_date = t.get("date", "")
        try:
            from datetime import datetime
            if isinstance(raw_date, (int, float)):
                date = datetime.fromtimestamp(raw_date).strftime("%Y-%m-%d")
            else:
                date = str(raw_date)[:10]
        except:
            date = str(raw_date)[:10]
        duration = t.get("duration", 0)
        duration_str = f"{duration // 60}min" if duration else "—"
        summary = t.get("summary", {})
        overview = summary.get("overview", "") if summary else ""
        action_items = summary.get("action_items", []) if summary else []
        if overview and len(overview) > 200:
            overview = overview[:200] + "..."
        lines.append(f"📅 *{title}*")
        lines.append(f"   Date: {date} | Duration: {duration_str}")
        if overview:
            lines.append(f"   {overview}")
        if action_items:
            clean_actions = [a.strip() for a in action_items if a and len(a.strip()) > 3]
            if clean_actions:
                lines.append(f"   Actions: {chr(10).join(f'   • {a[:100]}' for a in clean_actions[:3])}")
        lines.append("")
    return "\n".join(lines)

def get_transcript_detail(transcript_id):
    query = """
    query($id: String!) {
        transcript(id: $id) {
            title
            date
            duration
            summary {
                overview
                action_items
                keywords
            }
            sentences {
                text
                speaker_name
            }
        }
    }
    """
    try:
        data = graphql(query, {"id": transcript_id})
        return data.get("data", {}).get("transcript", {})
    except Exception as e:
        return {"error": str(e)}

def save_transcripts_to_memory(transcripts):
    """Save meeting summaries to Droplet memory for context."""
    from pathlib import Path
    memory_dir = Path("/root/.openclaw/workspace-hatfield/memory/meetings")
    memory_dir.mkdir(parents=True, exist_ok=True)
    saved = 0
    for t in transcripts:
        title = (t.get("title", "Untitled")).replace("/", "-")
        date = str(t.get("date", "unknown"))[:10]
        summary = t.get("summary", {}) or {}
        overview = summary.get("overview", "No summary available")
        action_items = summary.get("action_items", [])
        content = f"# {title}\nDate: {date}\n\n## Summary\n{overview}\n"
        if action_items:
            content += "\n## Action Items\n" + "\n".join(f"- {a}" for a in action_items)
        filename = f"{date}-{title[:50].replace(' ', '-')}.md"
        (memory_dir / filename).write_text(content)
        saved += 1
    return saved

def save_transcript_docx(transcript):
    """Save full transcript + summary to a .docx file. Returns the file path."""
    import re
    from pathlib import Path
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    title = transcript.get("title", "Untitled")
    raw_date = transcript.get("date", "")
    try:
        date_str = datetime.fromtimestamp(raw_date).strftime("%Y-%m-%d") if isinstance(raw_date, (int, float)) else str(raw_date)[:10]
    except Exception:
        date_str = str(raw_date)[:10]

    summary = transcript.get("summary", {}) or {}
    overview = summary.get("overview", "No summary available.")
    action_items = summary.get("action_items", []) or []
    sentences = transcript.get("sentences", []) or []

    doc = Document()

    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Date: {date_str}")
    doc.add_paragraph("")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(overview)

    if action_items:
        doc.add_heading("Action Items", level=1)
        for item in (action_items if isinstance(action_items, list) else action_items.strip().split("\n")):
            item = item.strip(" -•\n")
            if item:
                doc.add_paragraph(item, style="List Bullet")

    if sentences:
        doc.add_heading("Full Transcript", level=1)
        current_speaker = None
        current_para = None
        for s in sentences:
            speaker = s.get("speaker_name", "Unknown")
            text = s.get("text", "").strip()
            if not text:
                continue
            if speaker != current_speaker:
                current_para = doc.add_paragraph()
                run = current_para.add_run(f"{speaker}: ")
                run.bold = True
                current_para.add_run(text)
                current_speaker = speaker
            else:
                current_para.add_run(f" {text}")

    output_dir = Path("/opt/clawbot/transcripts")
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_title = re.sub(r"[^\w\s-]", "", title)[:50].strip().replace(" ", "-")
    filepath = output_dir / f"{date_str}-{safe_title}.docx"
    doc.save(str(filepath))
    audit("INFO", "fireflies-agent", f"Saved transcript docx: {filepath}")
    return str(filepath)


def send_transcript_email(filepath, transcript_title, date_str):
    """Email transcript docx from HatfieldFisher1013@gmail.com to nfisher@peak10group.com."""
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.base import MIMEBase
    from email import encoders

    GMAIL_USER = "HatfieldFisher1013@gmail.com"
    GMAIL_APP_PASSWORD = os.getenv("GMAIL_APP_PASSWORD")
    TO_EMAIL = "nfisher@peak10group.com"

    if not GMAIL_APP_PASSWORD:
        audit("ERROR", "fireflies-agent", "GMAIL_APP_PASSWORD not set — skipping email")
        return "⚠️ Email skipped: GMAIL_APP_PASSWORD not configured in vault"

    msg = MIMEMultipart()
    msg["From"] = GMAIL_USER
    msg["To"] = TO_EMAIL
    msg["Subject"] = f"Meeting Transcript: {transcript_title} ({date_str})"
    msg.attach(MIMEText(
        f"Attached is the full transcript and summary for:\n\n{transcript_title}\nDate: {date_str}",
        "plain"
    ))

    with open(filepath, "rb") as f:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(f.read())
    encoders.encode_base64(part)
    part.add_header("Content-Disposition", f'attachment; filename="{os.path.basename(filepath)}"')
    msg.attach(part)

    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(GMAIL_USER, GMAIL_APP_PASSWORD)
            server.sendmail(GMAIL_USER, TO_EMAIL, msg.as_string())
        audit("SUCCESS", "fireflies-agent", f"Email sent: {transcript_title}")
        return f"✅ Emailed to {TO_EMAIL}"
    except Exception as e:
        audit("ERROR", "fireflies-agent", f"Email failed: {e}")
        return f"❌ Email failed: {e}"


def run_transcript_pipeline(limit=5):
    """Full pipeline: fetch transcripts → save docx → email each one."""
    transcripts = get_recent_transcripts(limit)
    if not transcripts:
        return "No transcripts found in Fireflies."

    results = []
    for t in transcripts:
        title = t.get("title", "Untitled")
        raw_date = t.get("date", "")
        try:
            date_str = datetime.fromtimestamp(raw_date).strftime("%Y-%m-%d") if isinstance(raw_date, (int, float)) else str(raw_date)[:10]
        except Exception:
            date_str = str(raw_date)[:10]

        # Fetch full transcript with sentences
        detail = get_transcript_detail(t["id"])
        full_transcript = {**t, **(detail if detail and not detail.get("error") else {})}

        # Save docx
        filepath = save_transcript_docx(full_transcript)

        # Save to memory
        save_transcripts_to_memory([full_transcript])

        # Email
        email_result = send_transcript_email(filepath, title, date_str)

        results.append(f"📄 *{title}* ({date_str})\n   Saved: {filepath}\n   {email_result}")

    return "\n\n".join(results)


if __name__ == "__main__":
    print("Testing Fireflies pipeline...")
    print(run_transcript_pipeline(3))
