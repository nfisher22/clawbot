#!/usr/bin/env python3
"""
Fathom → OneDrive → Email Pipeline
====================================
1. Fetches recent meetings and transcripts from the Fathom API
2. Saves raw transcripts (.txt) to OneDrive: Fathom Meetings/Fathom Transcripts/
3. Generates a meeting summary DOCX via LLM and saves to OneDrive: Fathom Meetings/Fathom Summaries/
4. Emails the summary DOCX from nfisher@peak10group.com to nfisher@peak10group.com

Run once:   python3 fathom_onedrive_pipeline.py
Scheduled:  add to cron / systemd timer for continuous syncing
"""

import os
import io
import base64
import requests
import msal
from datetime import datetime, timezone
from pathlib import Path

from vault_secrets import get_secrets
from llm_client import chat_with_fallback

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

get_secrets()

# ── Configuration ──────────────────────────────────────────────────────────────
FATHOM_API_KEY      = os.getenv("FATHOM_API_KEY")
FATHOM_BASE         = "https://api.fathom.ai/external/v1"

AZURE_CLIENT_ID     = os.getenv("AZURE_CLIENT_ID")
AZURE_TENANT_ID     = os.getenv("AZURE_TENANT_ID")
AZURE_CLIENT_SECRET = os.getenv("AZURE_CLIENT_SECRET")
MS_USER_EMAIL       = os.getenv("MS_USER_EMAIL") or "nfisher@peak10group.com"

TRANSCRIPT_FOLDER   = "Fathom Meetings/Fathom Transcripts"
SUMMARY_FOLDER      = "Fathom Meetings/Fathom Summaries"
RECIPIENT_EMAIL     = "nfisher@peak10group.com"
AUDIT_LOG           = "/opt/clawbot/logs/audit.log"

# ── Audit Logging ──────────────────────────────────────────────────────────────
def audit(level, message):
    ts = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    line = f"{ts} | {level} | fathom-onedrive-pipeline | {message}\n"
    print(line.strip())
    try:
        with open(AUDIT_LOG, "a") as f:
            f.write(line)
    except Exception:
        pass


# ── Fathom API ─────────────────────────────────────────────────────────────────
def fathom_headers():
    return {"X-Api-Key": FATHOM_API_KEY, "Content-Type": "application/json"}


def get_fathom_meetings(limit=10):
    """Fetch recent meetings from Fathom including summaries and action items."""
    try:
        resp = requests.get(
            f"{FATHOM_BASE}/meetings",
            headers=fathom_headers(),
            params={
                "limit": limit,
                "include_summary": True,
                "include_action_items": True,
            },
            timeout=20,
        )
        if resp.status_code == 200:
            meetings = resp.json().get("items", [])
            audit("INFO", f"Fetched {len(meetings)} meetings from Fathom")
            return meetings
        audit("ERROR", f"Fathom meetings API {resp.status_code}: {resp.text[:200]}")
        return []
    except Exception as e:
        audit("ERROR", f"Fathom meetings exception: {e}")
        return []


def get_fathom_transcript(meeting_id):
    """
    Fetch the full transcript for a single meeting.
    Tries /meetings/{id}/transcript first, then /meetings/{id} with include_transcript param.
    Returns transcript text or None.
    """
    # Attempt 1: dedicated transcript endpoint
    try:
        resp = requests.get(
            f"{FATHOM_BASE}/meetings/{meeting_id}/transcript",
            headers=fathom_headers(),
            timeout=20,
        )
        if resp.status_code == 200:
            data = resp.json()
            # Try common response shapes
            text = (
                data.get("transcript")
                or data.get("text")
                or data.get("content")
                or _format_utterances(data.get("utterances") or data.get("entries") or [])
            )
            if text:
                return text
    except Exception:
        pass

    # Attempt 2: meeting detail with include_transcript param
    try:
        resp = requests.get(
            f"{FATHOM_BASE}/meetings/{meeting_id}",
            headers=fathom_headers(),
            params={"include_transcript": True},
            timeout=20,
        )
        if resp.status_code == 200:
            data = resp.json()
            text = (
                data.get("transcript")
                or data.get("transcript_text")
                or _format_utterances(data.get("utterances") or data.get("transcript_entries") or [])
            )
            if text:
                return text
    except Exception:
        pass

    return None


def _format_utterances(utterances):
    """Convert a list of speaker utterances to readable plain text."""
    if not utterances:
        return None
    lines = []
    for u in utterances:
        speaker = u.get("speaker") or u.get("speaker_name") or u.get("name") or "Speaker"
        content = u.get("text") or u.get("content") or u.get("transcript") or ""
        if content:
            lines.append(f"{speaker}: {content}")
    return "\n".join(lines) if lines else None


def build_transcript_text(meeting):
    """
    Build a best-effort transcript string from whatever Fathom provides.
    Falls back to the summary if no raw transcript is available.
    """
    meeting_id = meeting.get("id")

    # Try the API transcript endpoint first
    if meeting_id:
        transcript = get_fathom_transcript(meeting_id)
        if transcript:
            return transcript

    # Fall back to the summary fields embedded in the meeting object
    summary = meeting.get("default_summary") or {}
    overview = summary.get("markdown_formatted") or summary.get("text") or ""
    action_items = meeting.get("action_items") or []
    action_text = ""
    if action_items:
        action_text = "\n\nAction Items:\n" + "\n".join(
            f"- {a.get('description', '')}" for a in action_items if a.get("description")
        )
    return (overview + action_text).strip() or "(No transcript or summary available)"


# ── Microsoft Graph helpers ────────────────────────────────────────────────────
def get_graph_token():
    app = msal.ConfidentialClientApplication(
        AZURE_CLIENT_ID,
        authority=f"https://login.microsoftonline.com/{AZURE_TENANT_ID}",
        client_credential=AZURE_CLIENT_SECRET,
    )
    result = app.acquire_token_for_client(scopes=["https://graph.microsoft.com/.default"])
    return result.get("access_token")


def save_bytes_to_onedrive(token, content_bytes, content_type, folder, filename):
    """Upload raw bytes to OneDrive under folder/filename. Returns True on success."""
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": content_type,
    }
    url = (
        f"https://graph.microsoft.com/v1.0/users/{MS_USER_EMAIL}"
        f"/drive/root:/{folder}/{filename}:/content"
    )
    resp = requests.put(url, headers=headers, data=content_bytes, timeout=30)
    success = resp.status_code in (200, 201)
    level = "SUCCESS" if success else "ERROR"
    audit(level, f"OneDrive upload {'OK' if success else 'FAILED ' + str(resp.status_code)}: {folder}/{filename}")
    return success


def send_email_with_attachment(token, subject, body_html, attachment_name, attachment_bytes):
    """
    Send an email from MS_USER_EMAIL to RECIPIENT_EMAIL with a DOCX attachment
    via Microsoft Graph sendMail.
    """
    b64_content = base64.b64encode(attachment_bytes).decode("utf-8")
    payload = {
        "message": {
            "subject": subject,
            "body": {"contentType": "HTML", "content": body_html},
            "toRecipients": [{"emailAddress": {"address": RECIPIENT_EMAIL}}],
            "attachments": [
                {
                    "@odata.type": "#microsoft.graph.fileAttachment",
                    "name": attachment_name,
                    "contentType": (
                        "application/vnd.openxmlformats-officedocument"
                        ".wordprocessingml.document"
                    ),
                    "contentBytes": b64_content,
                }
            ],
        },
        "saveToSentItems": True,
    }
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json",
    }
    url = f"https://graph.microsoft.com/v1.0/users/{MS_USER_EMAIL}/sendMail"
    resp = requests.post(url, headers=headers, json=payload, timeout=30)
    success = resp.status_code == 202
    level = "SUCCESS" if success else "ERROR"
    audit(level, f"Email {'sent' if success else 'FAILED ' + str(resp.status_code)}: {subject}")
    return success


# ── LLM Summary ────────────────────────────────────────────────────────────────
def generate_summary(meeting_title, transcript_text):
    """Use the LLM to produce a concise, professional meeting summary."""
    prompt = f"""You are ClawBot, the AI Chief of Staff for Nathan Fisher at Peak 10 Group.
Summarize the following meeting transcript into a professional executive summary.

Structure your summary with these sections:
1. Meeting Overview (2-3 sentences)
2. Key Discussion Points (bullet list)
3. Decisions Made (bullet list, or "None recorded" if absent)
4. Action Items (bullet list with owner if known, or "None recorded" if absent)
5. Next Steps

Keep the tone professional and concise.

Meeting Title: {meeting_title}

Transcript / Notes:
{transcript_text[:6000]}
"""
    try:
        return chat_with_fallback([{"role": "user", "content": prompt}], max_tokens=1500)
    except Exception as e:
        audit("ERROR", f"LLM summary failed: {e}")
        return transcript_text  # fall back to raw text


# ── DOCX Generation ────────────────────────────────────────────────────────────
def _add_banner(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(4)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "1F3864")
    pPr.append(shd)
    run = para.add_run(f"  {text}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return para


def _add_body(doc, text, bold=False, font_size=10):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "Calibri"
    run.bold = bold
    return para


def build_summary_docx(meeting_title, meeting_date, duration_min, summary_text):
    """
    Build a polished DOCX meeting summary.
    Returns the document as bytes (BytesIO buffer).
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # Normal style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ── Title banner ──────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after  = Pt(4)
    pPr = title_para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "1F3864")
    pPr.append(shd)
    run = title_para.add_run(f"  Meeting Summary  |  {meeting_title}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xFF, 0xF2, 0xCC)

    # ── Metadata table ────────────────────────────────────────────────────────
    meta_table = doc.add_table(rows=3, cols=2)
    meta_data = [
        ("Date:",     meeting_date or "N/A"),
        ("Duration:", f"{duration_min} minutes" if duration_min else "N/A"),
        ("Prepared:", f"ClawBot — {datetime.now().strftime('%B %d, %Y')}"),
    ]
    for i, (label, value) in enumerate(meta_data):
        meta_table.cell(i, 0).text = label
        meta_table.cell(i, 1).text = value
        for cell in (meta_table.cell(i, 0), meta_table.cell(i, 1)):
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
        meta_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        meta_table.cell(i, 0).width = Inches(1.2)
        meta_table.cell(i, 1).width = Inches(5.3)

    doc.add_paragraph()  # spacer

    # ── Summary sections ──────────────────────────────────────────────────────
    _add_banner(doc, "EXECUTIVE SUMMARY")

    # Parse the LLM output into sections; if it's already structured, render as-is
    lines = summary_text.strip().splitlines()
    current_section = None
    buffer = []

    def flush_buffer():
        if buffer:
            for ln in buffer:
                stripped = ln.strip()
                if stripped.startswith(("-", "•", "*")):
                    para = doc.add_paragraph(style="List Bullet")
                    para.paragraph_format.space_after = Pt(1)
                    para.clear()
                    run = para.add_run(stripped.lstrip("-•* ").strip())
                    run.font.size = Pt(9.5)
                    run.font.name = "Calibri"
                elif stripped:
                    _add_body(doc, stripped)
            buffer.clear()

    section_keywords = {
        "meeting overview":   "MEETING OVERVIEW",
        "key discussion":     "KEY DISCUSSION POINTS",
        "decisions made":     "DECISIONS MADE",
        "action items":       "ACTION ITEMS",
        "next steps":         "NEXT STEPS",
    }

    for line in lines:
        matched_section = None
        for kw, banner in section_keywords.items():
            if kw in line.lower():
                matched_section = banner
                break

        if matched_section:
            flush_buffer()
            _add_banner(doc, matched_section)
            current_section = matched_section
        else:
            buffer.append(line)

    flush_buffer()

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f"Generated by ClawBot  |  Peak 10 Group  |  {datetime.now().strftime('%B %d, %Y')}"
    )
    footer_run.font.size  = Pt(8)
    footer_run.italic     = True
    footer_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Main Pipeline ──────────────────────────────────────────────────────────────
def run_pipeline(limit=10):
    audit("INFO", f"Starting Fathom → OneDrive pipeline (limit={limit})")

    token = get_graph_token()
    if not token:
        audit("ERROR", "Failed to acquire Microsoft Graph token — aborting")
        return

    meetings = get_fathom_meetings(limit)
    if not meetings:
        audit("INFO", "No meetings returned from Fathom — nothing to process")
        return

    processed = 0
    for meeting in meetings:
        meeting_id    = meeting.get("id", "unknown")
        title         = (
            meeting.get("title")
            or meeting.get("meeting_title")
            or "Untitled Meeting"
        )
        raw_date      = meeting.get("created_at") or meeting.get("scheduled_start_time") or ""
        date_str      = raw_date[:10] if raw_date else datetime.now().strftime("%Y-%m-%d")
        duration_min  = meeting.get("duration_seconds", 0) // 60 if meeting.get("duration_seconds") else None
        safe_title    = title[:60].replace("/", "-").replace("\\", "-").replace(":", "-").strip()
        base_filename = f"{date_str}-{safe_title}"

        audit("INFO", f"Processing: {title} ({date_str})")

        # ── Step 1: Get transcript ─────────────────────────────────────────
        transcript_text = build_transcript_text(meeting)

        # ── Step 2: Save raw transcript to OneDrive ────────────────────────
        transcript_filename = f"{base_filename}.txt"
        save_bytes_to_onedrive(
            token,
            transcript_text.encode("utf-8"),
            "text/plain",
            TRANSCRIPT_FOLDER,
            transcript_filename,
        )

        # ── Step 3: Generate LLM summary ──────────────────────────────────
        audit("INFO", f"Generating summary for: {title}")
        summary_text = generate_summary(title, transcript_text)

        # ── Step 4: Build DOCX ─────────────────────────────────────────────
        docx_bytes = build_summary_docx(title, date_str, duration_min, summary_text)
        docx_filename = f"{base_filename}-Summary.docx"

        # ── Step 5: Save summary DOCX to OneDrive ─────────────────────────
        save_bytes_to_onedrive(
            token,
            docx_bytes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            SUMMARY_FOLDER,
            docx_filename,
        )

        # ── Step 6: Email the summary ──────────────────────────────────────
        email_subject = f"Meeting Summary: {title} — {date_str}"
        email_body = f"""<div style="font-family:Calibri,sans-serif;font-size:14px;">
<p>Hi Nathan,</p>
<p>Please find attached the meeting summary for <strong>{title}</strong> ({date_str}).</p>
<p>The summary has also been saved to your OneDrive under
<strong>Fathom Meetings &rsaquo; Fathom Summaries</strong>.</p>
<pre style="background:#f4f4f4;padding:12px;border-radius:4px;font-size:13px;white-space:pre-wrap;">{summary_text[:1500]}{"..." if len(summary_text) > 1500 else ""}</pre>
<br>
<p>Thanks,</p>
<p><strong>ClawBot</strong><br>AI Chief of Staff<br>Peak 10 Group</p>
</div>"""

        send_email_with_attachment(
            token,
            email_subject,
            email_body,
            docx_filename,
            docx_bytes,
        )

        processed += 1

    audit("INFO", f"Pipeline complete — processed {processed}/{len(meetings)} meetings")
    print(f"\n✅ Done. Processed {processed} meeting(s).")
    print(f"   Transcripts → OneDrive: {TRANSCRIPT_FOLDER}/")
    print(f"   Summaries   → OneDrive: {SUMMARY_FOLDER}/")
    print(f"   Emails sent → {RECIPIENT_EMAIL}")


if __name__ == "__main__":
    import sys
    limit = int(sys.argv[1]) if len(sys.argv) > 1 else 10
    run_pipeline(limit=limit)
